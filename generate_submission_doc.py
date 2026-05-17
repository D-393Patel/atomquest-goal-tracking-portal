from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "AtomQuest_Hackathon_Submission.docx"
DIAGRAM_PATH = ROOT / "architecture.png"
REPO_URL = "https://github.com/D-393Patel/atomquest-goal-tracking-portal.git"
DEPLOYED_URL = "https://d-393patel.github.io/atomquest-goal-tracking-portal/"


def font(size, bold=False):
    try:
        return ImageFont.truetype("arialbd.ttf" if bold else "arial.ttf", size)
    except Exception:
        return ImageFont.load_default()


def draw_center(draw, box, text, fill="#123233", size=24, bold=False):
    fnt = font(size, bold)
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + (len(lines) - 1) * 8
    y = box[1] + ((box[3] - box[1]) - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        x = box[0] + ((box[2] - box[0]) - w) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += h + 8


def make_diagram():
    if Image is None:
        return None

    img = Image.new("RGB", (1400, 850), "#eef6f5")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((70, 95, 1330, 760), radius=24, fill="#ffffff", outline="#c9d8d6", width=4)
    draw_center(draw, (70, 20, 1330, 90), "AtomQuest Portal Architecture", size=40, bold=True)

    boxes = [
        ((120, 245, 390, 365), "#d9eeee", "#126c67", "User Browser\nEmployee / Manager / Admin"),
        ((555, 170, 845, 295), "#e8edf3", "#607085", "Web Portal UI\nDashboards, forms, role views"),
        ((555, 410, 845, 535), "#fff0cf", "#b98319", "Business Logic\nValidation, scoring, locking"),
        ((1010, 255, 1280, 380), "#e1edf9", "#306fb2", "Data Store\nGoals, users, cycles"),
        ((1010, 500, 1280, 610), "#fce1e5", "#b74253", "Audit Logs\nWho changed what and when"),
        ((155, 500, 425, 610), "#dff4e8", "#178454", "Reports\nAchievement CSV export"),
    ]

    for box, fill, outline, text in boxes:
        draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=4)
        title, subtitle = text.split("\n", 1)
        draw_center(draw, (box[0], box[1] + 12, box[2], box[1] + 62), title, size=27, bold=True)
        draw_center(draw, (box[0] + 12, box[1] + 62, box[2] - 12, box[3] - 10), subtitle, size=20)

    arrows = [
        ((390, 305), (555, 235)),
        ((700, 295), (700, 410)),
        ((845, 475), (1010, 320)),
        ((845, 505), (1010, 555)),
        ((555, 500), (425, 555)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#49636a", width=5)
        ex, ey = end
        sx, sy = start
        if ex > sx:
            points = [(ex, ey), (ex - 18, ey - 10), (ex - 12, ey + 13)]
        elif ex < sx:
            points = [(ex, ey), (ex + 18, ey - 10), (ex + 12, ey + 13)]
        else:
            points = [(ex, ey), (ex - 12, ey - 18), (ex + 12, ey - 18)]
        draw.polygon(points, fill="#49636a")

    img.save(DIAGRAM_PATH)
    return DIAGRAM_PATH


def set_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(18, 108, 103)


def main():
    diagram = make_diagram()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AtomQuest Hackathon 1.0 Submission")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 55, 56)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("In-House Goal Setting & Tracking Portal")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(96, 112, 133)

    doc.add_paragraph()

    add_heading(doc, "Submission Links")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    data = [
        ("Working Link", DEPLOYED_URL),
        ("Source Code Repository", REPO_URL),
        ("Architecture Diagram", "Embedded below in this document; source file: architecture.svg"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell(row.cells[0], label, bold=True)
        set_cell(row.cells[1], value)

    add_heading(doc, "Demo Credentials")
    creds = doc.add_table(rows=4, cols=3)
    creds.style = "Table Grid"
    headers = ["Role", "Email", "Password"]
    rows = [
        ["Employee", "employee@test.com", "123456"],
        ["Manager", "manager@test.com", "123456"],
        ["Admin / HR", "admin@test.com", "123456"],
    ]
    for i, header in enumerate(headers):
        set_cell(creds.rows[0].cells[i], header, bold=True)
    for r, row_data in enumerate(rows, start=1):
        for c, value in enumerate(row_data):
            set_cell(creds.rows[r].cells[c], value)

    add_heading(doc, "Solution Summary")
    summary = doc.add_paragraph()
    summary.add_run(
        "AtomQuest Goal Setting & Tracking Portal is a browser-based demo solution for employee goal creation, manager approval, quarterly achievement tracking, check-ins, admin governance, audit logs, analytics, shared goals, and CSV reporting."
    )

    add_heading(doc, "Implemented Scope")
    features = [
        "Employee goal creation with total 100% weightage validation, minimum 10% per goal, and maximum 8 goals.",
        "Manager approval workflow with inline target and weightage editing, approval, and return for rework.",
        "Goal locking after approval and Admin unlock exception handling.",
        "Shared departmental goals pushed to multiple employees with title/target read-only and weightage adjustable.",
        "Quarterly achievement updates with UoM-based progress score calculation.",
        "Manager check-in comments, completion dashboard, audit trail, cycle controls, analytics, escalation panel, and CSV export.",
    ]
    for feature in features:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(feature)

    add_heading(doc, "Recommended Judge Demo Flow")
    flow = [
        "Login as Employee and show goal sheet, validations, and quarterly update screen.",
        "Switch to Manager and show approvals, inline edits, approval/return, shared goals, and check-in comments.",
        "Switch to Admin and show completion dashboard, cycle management, audit logs, unlock action, analytics, and CSV export.",
    ]
    for item in flow:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

    add_heading(doc, "Architecture Diagram")
    if diagram:
        doc.add_picture(str(diagram), width=Inches(6.8))
    else:
        doc.add_paragraph("Architecture diagram source is included in the project as architecture.svg.")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
